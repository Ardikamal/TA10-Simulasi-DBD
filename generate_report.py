# generate_report.py 
from simulasidbd_core import run_pipeline
from docx import Document
import sys

def make_docx(results, outname="Laporan_TA10_DBD.docx"):
    df = results["df"]
    doc = Document()
    doc.add_heading("Laporan TA-10 — Simulasi DBD Jakarta", level=1)

    doc.add_heading("Hasil Model", level=2)
    for key in ["sir", "rm"]:
        if key in results:
            m = results[key]
            doc.add_paragraph(f"{key.upper()} → RMSE = {m['rmse']:.3f}")

    doc.add_heading("Preview Data", level=2)
    table = doc.add_table(rows=1, cols=len(df.columns))
    hdr = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr[i].text = col

    for i in range(min(10, len(df))):
        row = table.add_row().cells
        for j, col in enumerate(df.columns):
            row[j].text = str(df.iloc[i][col])

    doc.save(outname)

if __name__ == "__main__":
    csv = sys.argv[1] if len(sys.argv) > 1 else "DATA DBD.csv"
    res = run_pipeline(csv)
    make_docx(res)
