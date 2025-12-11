# app.py
import streamlit as st
import pandas as pd
import matplotlib.pyplot as plt
from simulasidbd_core import run_pipeline, load_and_prep
import numpy as np
import io
from docx import Document
import zipfile
import os

# ============================================================
# Page Setup
# ============================================================
st.set_page_config(
    page_title="Simulasi DBD TA-10",
    layout="wide",
    page_icon="🦟"
)

# ============================================================
# Glassmorphism CSS
# ============================================================
st.markdown("""
<style>
body {
    background: linear-gradient(135deg, #d6eaff 0%, #f6e6ff 50%, #ffe7e7 100%) !important;
}
.block-container {
    padding-top: 2rem;
}
.glass-card {
    background: rgba(255, 255, 255, 0.18);
    border-radius: 16px;
    padding: 25px;
    border: 1px solid rgba(255, 255, 255, 0.35);
    backdrop-filter: blur(14px) saturate(180%);
    -webkit-backdrop-filter: blur(14px) saturate(180%);
    box-shadow: 0 8px 25px rgba(0,0,0,0.15);
    margin-bottom: 20px;
}
.metric-card {
    background: rgba(255, 255, 255, 0.25);
    padding: 18px;
    text-align: center;
    border-radius: 15px;
    border: 1px solid rgba(255,255,255,0.3);
    backdrop-filter: blur(10px);
    margin-bottom: 12px;
}
.stButton>button {
    background: rgba(255, 255, 255, 0.2) !important;
    color: #000 !important;
    backdrop-filter: blur(6px);
    border-radius: 12px !important;
    padding: 10px 18px !important;
    border: 1px solid rgba(255,255,255,0.4) !important;
    font-weight: 600;
}
.stButton>button:hover {
    background: rgba(255, 255, 255, 0.35) !important;
}
h1, h2, h3 { font-family: 'Segoe UI', sans-serif !important; }
</style>
""", unsafe_allow_html=True)

# ============================================================
# Header
# ============================================================
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.title("🧊 Simulasi Penyebaran DBD — TA-10 (Ardi Kamal Karima)")
st.write("Analisis epidemiologi menggunakan model SIR & Ross–Macdonald, lengkap dengan prediksi dan interpretasi otomatis.")
st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# Upload Section
# ============================================================
st.markdown('<div class="glass-card">', unsafe_allow_html=True)
st.header("📂 Upload Dataset CSV")

uploaded = st.file_uploader("Upload file CSV kasus DBD", type=["csv"])

if uploaded:
    df = pd.read_csv(uploaded)
    csv_path = "uploaded_dataset.csv"
    df.to_csv(csv_path, index=False)
elif os.path.exists("DATA DBD.csv"):
    csv_path = "DATA DBD.csv"
    df = pd.read_csv(csv_path)
else:
    st.info("Silakan upload dataset terlebih dahulu atau letakkan file DATA DBD.csv.")
    st.stop()

st.subheader("📌 Preview Dataset")
st.dataframe(df.head(), use_container_width=True)
st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# Layout 2 Columns
# ============================================================
left, right = st.columns([1.1, 1.6], gap="large")

# ============================================================
# LEFT PANEL — SETTINGS
# ============================================================
with left:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("⚙️ Pengaturan Model")

    use_sir = st.checkbox("Gunakan Model SIR", value=True)
    use_rm = st.checkbox("Gunakan Model Ross–Macdonald", value=True)

    st.write("### 📅 Kolom Dataset")
    date_col = st.selectbox("Kolom Tanggal:", df.columns)
    case_col = st.selectbox("Kolom Kasus:", df.columns)
    use_manual = st.checkbox("Gunakan Kolom Manual", value=False)

    st.write("### 🧮 Parameter Model")
    N_h = st.number_input("Populasi Manusia (N_h)", value=5000000)
    N_v = st.number_input("Populasi Nyamuk (N_v)", value=500000)
    Iv0 = st.number_input("Nyamuk Terinfeksi Awal (I_v0)", value=1000)

    start = st.button("🚀 Jalankan Simulasi", type="primary")
    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# RIGHT PANEL — RESULTS
# ============================================================
with right:

    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("📊 Hasil Simulasi")

    if start:
        with st.spinner("Menjalankan simulasi..."):
            if use_manual:
                results = run_pipeline(csv_path, use_sir, use_rm, N_h, N_v, Iv0,
                                       manual_date_col=date_col,
                                       manual_case_col=case_col)
            else:
                results = run_pipeline(csv_path, use_sir, use_rm, N_h, N_v, Iv0)

        df = results["df"]
        t = results["t"]
        I = results["I"]

        # ==============================
        # Grafik Simulasi
        # ==============================
        fig, ax = plt.subplots(figsize=(12,5))
        ax.scatter(t, I, color="black", label="Data Asli", s=20)

        metrics = []

        if "sir" in results:
            sol = results["sir"]["sol"]
            rmse = results["sir"]["rmse"]
            ax.plot(t, sol[:,1], label=f"SIR (RMSE={rmse:.2f})", linewidth=2)
            metrics.append(("Model SIR", rmse))

        if "rm" in results:
            sol = results["rm"]["sol"]
            rmse = results["rm"]["rmse"]
            ax.plot(t, sol[:,0], label=f"Ross–Macdonald (RMSE={rmse:.2f})", linewidth=2)
            metrics.append(("Ross–Macdonald", rmse))

        ax.legend()
        ax.grid()
        st.pyplot(fig)

    st.markdown('</div>', unsafe_allow_html=True)

# ==============================
# METRIC CARDS
# ==============================
if start:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("📌 Ringkasan Nilai RMSE")

    for name, score in metrics:
        st.markdown(f"""
        <div class="metric-card">
            <h3>{name}</h3>
            <h2>{score:.3f}</h2>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# INTERPRETASI HASIL
# ============================================================
if start:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("📘 Interpretasi Hasil Simulasi")

    if "sir" in results:
        sir_last = results["sir"]["sol"][-1,1]
        sir_prev = results["sir"]["sol"][-5,1]
        sir_trend = "meningkat" if sir_last > sir_prev else "menurun"

        st.markdown(f"""
        ### 🔷 Model SIR
        • Nilai kasus simulasi terakhir = **{sir_last:.2f}**  
        • Tren 5 hari terakhir: **{sir_trend}**  
        """)

    if "rm" in results:
        rm_last = results["rm"]["sol"][-1,0]
        rm_prev = results["rm"]["sol"][-5,0]
        rm_trend = "meningkat" if rm_last > rm_prev else "menurun"

        st.markdown(f"""
        ### 🔶 Model Ross–Macdonald
        • Prediksi populasi manusia terinfeksi (Ih) = **{rm_last:.2f}**  
        • Tren 5 hari terakhir: **{rm_trend}**  
        """)

    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# 7-DAY FORECAST
# ============================================================
if start:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("🔮 Prediksi 7 Hari Ke Depan")

    forecast_days = 7
    t_future = np.arange(len(t), len(t) + forecast_days)

    pred_sir, pred_rm = None, None

    # SIR forecast
    if "sir" in results:
        β, γ = results["sir"]["params"]
        S0, I0, R0 = results["sir"]["sol"][-1]
        pred = []
        S, I_sim, R = S0, I0, R0
        for _ in range(forecast_days):
            dS = -β * S * I_sim
            dI = β * S * I_sim - γ * I_sim
            dR = γ * I_sim
            S += dS
            I_sim += dI
            R += dR
            pred.append(I_sim)
        pred_sir = np.array(pred)

    # Ross–Macdonald forecast
    if "rm" in results:
        a, b, c, μv = results["rm"]["params"]
        Sh, Ih, Rh, Sv, Iv = results["rm"]["sol"][-1]
        pred = []
        for _ in range(forecast_days):
            dSh = -a*b*Sh*Iv
            dIh = a*b*Sh*Iv - 0.1*Ih
            dRh = 0.1*Ih
            dSv = -a*c*Sv*Ih - μv*Sv
            dIv = a*c*Sv*Ih - μv*Iv
            Sh += dSh
            Ih += dIh
            Rh += dRh
            Sv += dSv
            Iv += dIv
            pred.append(Ih)
        pred_rm = np.array(pred)

    forecast_df = pd.DataFrame({
        "Hari ke-": np.arange(1, forecast_days + 1),
        "Prediksi SIR": pred_sir,
        "Prediksi RM (Ih)": pred_rm
    })
    st.dataframe(forecast_df, use_container_width=True)

    fig2, ax2 = plt.subplots(figsize=(12,5))
    ax2.scatter(t, I, color="black", label="Data Asli")

    if pred_sir is not None:
        ax2.plot(t_future, pred_sir, "--", label="Prediksi SIR 7 Hari")

    if pred_rm is not None:
        ax2.plot(t_future, pred_rm, "--", label="Prediksi RM 7 Hari")

    ax2.legend()
    ax2.grid()
    st.pyplot(fig2)

    st.markdown("### 🧠 Interpretasi Prediksi")
    if pred_sir is not None:
        trend_sir = "naik" if pred_sir[-1] > pred_sir[0] else "turun"
        st.write(f"Model **SIR** memprediksi tren **{trend_sir}** dalam 7 hari ke depan.")

    if pred_rm is not None:
        trend_rm = "naik" if pred_rm[-1] > pred_rm[0] else "turun"
        st.write(f"Model **Ross–Macdonald** memprediksi jumlah infeksi manusia akan **{trend_rm}**.")

    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# PARAMETER INTERPRETATION
# ============================================================
if start:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("🧬 Interpretasi Parameter Model")

    # SIR Params
    if "sir" in results:
        β, γ = results["sir"]["params"]
        st.markdown(f"""
        ### 🔷 Parameter Model SIR  
        **β (beta) = {β:.4f}**  
        → Laju penularan.  
        
        **γ (gamma) = {γ:.4f}**  
        → Laju kesembuhan.  
        → Rata-rata lama sakit = **{1/γ:.2f} hari**  
        """)

    # RM Params
    if "rm" in results:
        a, b, c, μv = results["rm"]["params"]
        st.markdown(f"""
        ### 🔶 Parameter Model Ross–Macdonald  
        **a = {a:.4f}** → Frekuensi gigitan nyamuk per hari  
        **b = {b:.4f}** → Probabilitas manusia tertular  
        **c = {c:.4f}** → Probabilitas nyamuk tertular  
        **μv = {μv:.4f}** → Laju kematian nyamuk  
        → Umur nyamuk ≈ **{1/μv:.2f} hari**  
        """)

    st.markdown('</div>', unsafe_allow_html=True)

# ============================================================
# DOWNLOAD REPORTS
# ============================================================
if start:
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    st.subheader("📁 Download Hasil")

    # DOCX
    if st.button("📄 Generate Laporan (.docx)"):
        doc = Document()
        doc.add_heading("Laporan TA-10 — Simulasi DBD", level=1)

        for m, v in metrics:
            doc.add_paragraph(f"{m}: RMSE = {v:.3f}")

        doc.save("Laporan_TA10.docx")
        with open("Laporan_TA10.docx", "rb") as f:
            st.download_button("Download Laporan", f, file_name="Laporan_TA10.docx")

    # ZIP
    if st.button("📦 Download Semua Data (.zip)"):

        buffer = io.BytesIO()
        with zipfile.ZipFile(buffer, "w") as zf:
            zf.writestr("summary.txt", "\n".join([f"{m}: {v}" for m, v in metrics]))
            zf.writestr("dataset.csv", df.to_csv(index=False))
        buffer.seek(0)

        st.download_button("Download ZIP", buffer, "TA10_Hasil.zip")

    st.markdown('</div>', unsafe_allow_html=True)
